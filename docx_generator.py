"""
Spieler-Etiketten Generator
Alle Maße 1:1 aus der Vorlage (260224_FC Hansa Rostock.docx).

Foto:      568800 × 737999 EMU  = 1,58 × 2,05 cm  (beide Achsen exakt)
Zeilenhöhen: 249 / 249 / 1179 / 249 DXA  (exact)
Seitenbreite: 20 cm Nutzbreite, 0,5 cm Rand links/rechts  (wie Original)

Seitenumbruch-Schutz:
  Vor jedem Paar wird geprüft, ob es noch auf die Seite passt.
  Falls nicht → Seitenumbruch vor dem Paar. Kein Paar wird je gespalten.
"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── Exakte Original-Maße ────────────────────────────────────────────────────

# Gitterspalten pro Spieler in DXA (Twips); 16 Spalten für beide Spieler
_GRID = [680, 454, 227, 453, 1134, 680, 681, 1361]
_GRID_FULL = _GRID + _GRID
_PAGE_CONTENT_DXA = sum(_GRID_FULL)         # 11340 DXA = 20,00 cm

# Zeilenhöhen (DXA, exact) — 1:1 aus Original
_H = [249, 249, 1179, 249]                  # Kopf / Stats / Foto / Transfer
_PAIR_HEIGHT_DXA = sum(_H)                  # 1926 DXA = 3,40 cm

# Foto-Dimensionen (EMU) — 1:1 aus Original-XML
_PHOTO_W_EMU = 568800                        # 1,58 cm
_PHOTO_H_EMU = 737999                        # 2,05 cm

# Seiten-Setup
_PAGE_H_DXA   = int(29.7 / 2.54 * 1440)    # A4 Höhe ≈ 16776 DXA
_MARGIN_T_DXA = int(1.50 / 2.54 * 1440)    # oben  ≈  849 DXA
_MARGIN_B_DXA = int(1.50 / 2.54 * 1440)    # unten ≈  849 DXA
_TITLE_DXA    = 420                          # ca. Titel + Abstand
_USABLE_FIRST = _PAGE_H_DXA - _MARGIN_T_DXA - _MARGIN_B_DXA - _TITLE_DXA
_USABLE_REST  = _PAGE_H_DXA - _MARGIN_T_DXA - _MARGIN_B_DXA


# ─── Zellen-Merge-Definitionen (start, end), offset je Spieler = 0 oder 8 ────
# Kopfzeile
_H_NUM  = (0, 2)   # 1361 DXA
_H_NAME = (3, 6)   # 2948 DXA
_H_POS  = (7, 7)   # 1361 DXA
# Stats-Zeile
_S_INF  = (0, 3)   # 1814 DXA
_S_GAM  = (4, 4)   # 1134 DXA
_S_GOA  = (5, 5)   #  680 DXA
_S_ASS  = (6, 6)   #  681 DXA
_S_CAR  = (7, 7)   # 1361 DXA
# Foto-Zeile
_F_PHO  = (0, 2)   # 1361 DXA
_F_NOT  = (3, 7)   # 4309 DXA
# Transfer-Zeile
_T_ALL  = (0, 7)   # 5670 DXA


# ─── Hilfsfunktionen ─────────────────────────────────────────────────────────

def _sum(span):
    return sum(_GRID[span[0]:span[1]+1])


def _set_grid(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Gesamtbreite
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), str(_PAGE_CONTENT_DXA))
    tw.set(qn('w:type'), 'dxa')
    tblPr.append(tw)
    # tblGrid
    old = tbl.find(qn('w:tblGrid'))
    if old is not None:
        tbl.remove(old)
    tg = OxmlElement('w:tblGrid')
    for w in _GRID_FULL:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tg.append(gc)
    tblPr.addnext(tg)


def _row_height(row, dxa):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    el = OxmlElement('w:trHeight')
    el.set(qn('w:val'), str(dxa))
    el.set(qn('w:hRule'), 'exact')
    trPr.append(el)
    # cantSplit: Zeile wird nie durch Seitenumbruch geteilt
    for old in trPr.findall(qn('w:cantSplit')):
        trPr.remove(old)
    trPr.append(OxmlElement('w:cantSplit'))


def _cell_w(cell, dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    el = OxmlElement('w:tcW')
    el.set(qn('w:w'), str(dxa))
    el.set(qn('w:type'), 'dxa')
    tcPr.append(el)


def _valign(cell, v='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    el = OxmlElement('w:vAlign')
    el.set(qn('w:val'), v)
    tcPr.append(el)


def _text(cell, txt, bold=False, sz=9,
          align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=False):
    """Text in Zelle schreiben; keep_next=True fügt keepWithNext hinzu."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if keep_next:
        p.paragraph_format.keep_with_next = True
    if txt:
        r = p.add_run(txt)
        r.bold = bold
        r.font.size = Pt(sz)


def _photo(cell, path, keep_next=False):
    """Foto in Zelle, exakt 568800×737999 EMU wie im Original."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if keep_next:
        p.paragraph_format.keep_with_next = True
    try:
        run = p.add_run()
        run.add_picture(path, width=Emu(_PHOTO_W_EMU), height=Emu(_PHOTO_H_EMU))
    except Exception:
        p.add_run('[Foto]').font.size = Pt(7)


def _merge(table, row, s, e):
    if s == e:
        return table.cell(row, s)
    return table.cell(row, s).merge(table.cell(row, e))


def shorten_pos(pos):
    m = {
        'Torwart':                  'TW',
        'Innenverteidiger':         'IV',
        'Rechter Verteidiger':      'RV',
        'Linker Verteidiger':       'LV',
        'Rechter Abwehrspieler':    'RAV',
        'Linker Abwehrspieler':     'LAV',
        'Defensives Mittelfeld':    'ZDM',
        'Zentrales Mittelfeld':     'ZM',
        'Rechtes Mittelfeld':       'RM',
        'Linkes Mittelfeld':        'LM',
        'Offensives Mittelfeld':    'OM',
        'Rechter Flügel':           'RF',
        'Linker Flügel':            'LF',
        'Linksaußen':               'LA',
        'Rechtsaußen':              'RA',
        'Hängende Spitze':          'HS',
        'Mittelstürmer':            'ST',
        'Sturm':                    'ST',
    }
    return m.get(pos, pos[:4] if len(pos) > 5 else pos)


# Hintergrundfarben je Positionsgruppe (hex, ohne #)
_POS_COLORS = {
    # Torwart – Gold/Gelb
    'TW':  'FFD966',
    # Abwehr – Blau
    'IV':  '9DC3E6',
    'LV':  '9DC3E6',
    'RV':  '9DC3E6',
    'LAV': '9DC3E6',
    'RAV': '9DC3E6',
    # Mittelfeld – Grün
    'ZDM': 'A9D18E',
    'ZM':  'A9D18E',
    'DM':  'A9D18E',
    'OM':  'A9D18E',
    'LM':  'A9D18E',
    'RM':  'A9D18E',
    'LF':  'A9D18E',
    'RF':  'A9D18E',
    # Angriff – Orange
    'LA':  'F4B183',
    'RA':  'F4B183',
    'ST':  'F4B183',
    'HS':  'F4B183',
    'MS':  'F4B183',
}
_COLOR_DEFAULT = 'F2F2F2'   # Hellgrau für unbekannte Positionen


def _pos_color(pos_raw):
    """Gibt Hex-Farbe für eine Position zurück."""
    short = shorten_pos(pos_raw)
    return _POS_COLORS.get(short, _COLOR_DEFAULT)


def _cell_bg(cell, hex_color):
    """Hintergrundfarbe einer Zelle setzen (hex ohne #, z.B. 'FFD966')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Altes shd entfernen
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


# ─── Ein Spielerpaar als Tabelle ─────────────────────────────────────────────

def _add_pair_table(doc, p1, p2):
    """Erstellt eine 4-Zeilen-Tabelle mit exakten Original-Maßen."""

    cards = lambda p: f"{p.get('yellow','0')}/{p.get('yellow_red','0')}/{p.get('red','0')}"
    info  = lambda p: f"{p.get('age','-')} / {p.get('height','-')} / {p.get('nationality','-')}"

    tbl = doc.add_table(rows=4, cols=16)
    tbl.style = 'Table Grid'
    _set_grid(tbl)

    for ri, h in enumerate(_H):
        _row_height(tbl.rows[ri], h)

    # ── Zeile 0: Kopf ──────────────────────────────────────────────────────
    for off, pl in [(0, p1), (8, p2)]:
        num  = _merge(tbl, 0, off+_H_NUM[0],  off+_H_NUM[1])
        name = _merge(tbl, 0, off+_H_NAME[0], off+_H_NAME[1])
        pos  = _merge(tbl, 0, off+_H_POS[0],  off+_H_POS[1])
        _cell_w(num, _sum(_H_NUM)); _cell_w(name, _sum(_H_NAME)); _cell_w(pos, _sum(_H_POS))
        _valign(num); _valign(name); _valign(pos)
        # Kopfzeile farbig nach Positionsgruppe
        color = _pos_color(pl.get('position', ''))
        _cell_bg(num, color); _cell_bg(name, color); _cell_bg(pos, color)
        _text(num,  str(pl.get('number','')),                           bold=True, sz=10, keep_next=True)
        _text(name, f"{pl.get('name','')} ({pl.get('foot','-')})",      bold=True, sz=9,  keep_next=True)
        _text(pos,  shorten_pos(pl.get('position','')),                  bold=True, sz=9,  keep_next=True)

    # ── Zeile 1: Stats ─────────────────────────────────────────────────────
    for off, pl in [(0, p1), (8, p2)]:
        ci = _merge(tbl, 1, off+_S_INF[0], off+_S_INF[1])
        cg = _merge(tbl, 1, off+_S_GAM[0], off+_S_GAM[1])
        co = _merge(tbl, 1, off+_S_GOA[0], off+_S_GOA[1])
        ca = _merge(tbl, 1, off+_S_ASS[0], off+_S_ASS[1])
        cc = _merge(tbl, 1, off+_S_CAR[0], off+_S_CAR[1])
        for c, w in [(ci,_sum(_S_INF)),(cg,_sum(_S_GAM)),(co,_sum(_S_GOA)),(ca,_sum(_S_ASS)),(cc,_sum(_S_CAR))]:
            _cell_w(c, w); _valign(c)
        is_gk = pl.get('position', '').strip() == 'Torwart'
        _text(ci, info(pl),                    sz=8, keep_next=True)
        _text(cg, pl.get('games_display', '-'), sz=8, keep_next=True)
        _text(co, pl.get('gegentore', '-') if is_gk else pl.get('goals', '0'),  sz=8, keep_next=True)
        _text(ca, pl.get('zu_null', '-')   if is_gk else pl.get('assists', '0'), sz=8, keep_next=True)
        _text(cc, cards(pl),                   sz=8, keep_next=True)

    # ── Zeile 2: Foto + Notiz ───────────────────────────────────────────────
    for off, pl in [(0, p1), (8, p2)]:
        ph = _merge(tbl, 2, off+_F_PHO[0], off+_F_PHO[1])
        nt = _merge(tbl, 2, off+_F_NOT[0], off+_F_NOT[1])
        _cell_w(ph, _sum(_F_PHO)); _cell_w(nt, _sum(_F_NOT))
        _valign(ph, 'center'); _valign(nt, 'top')
        if pl.get('local_photo') and os.path.exists(pl['local_photo']):
            _photo(ph, pl['local_photo'], keep_next=True)
        else:
            _text(ph, '[Foto]', sz=7, keep_next=True)
        _text(nt, '', sz=8, keep_next=True)

    # ── Zeile 3: Transfer ───────────────────────────────────────────────────
    for off, pl in [(0, p1), (8, p2)]:
        ct = _merge(tbl, 3, off+_T_ALL[0], off+_T_ALL[1])
        _cell_w(ct, _sum(_T_ALL)); _valign(ct)
        _text(ct, pl.get('transfer_info','-'), sz=7.5,
              align=WD_ALIGN_PARAGRAPH.CENTER)


# ─── Dokument erstellen ──────────────────────────────────────────────────────

def generate_etiketten(players, output_path, team_name=''):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(0.50)
    sec.right_margin  = Cm(0.50)
    sec.top_margin    = Cm(1.50)
    sec.bottom_margin = Cm(1.50)

    # Titel
    if team_name:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(4)
        r = tp.add_run(f'Spieleretiketten – {team_name}')
        r.bold = True
        r.font.size = Pt(13)

    # Seitenkapazität tracken
    used_dxa   = _TITLE_DXA   # erste Seite hat Titelplatz belegt
    usable_dxa = _USABLE_FIRST

    _POS_ORDER = {
        'Torwart': 0,
        'Innenverteidiger': 1, 'Rechter Verteidiger': 1, 'Linker Verteidiger': 1,
        'Defensives Mittelfeld': 2, 'Zentrales Mittelfeld': 2,
        'Rechtes Mittelfeld': 2, 'Linkes Mittelfeld': 2, 'Offensives Mittelfeld': 2,
        'Rechter Flügel': 2, 'Linker Flügel': 2, 'Linksaußen': 2, 'Rechtsaußen': 2,
        'Mittelstürmer': 3, 'Hängende Spitze': 3, 'Sturm': 3,
    }
    players = sorted(players, key=lambda p: (
        _POS_ORDER.get(p.get('position', ''), 2),
        int(p.get('number', 99)) if str(p.get('number', '')).isdigit() else 99
    ))

    pairs = []
    for i in range(0, len(players), 2):
        pairs.append((players[i], players[i+1] if i+1 < len(players) else _empty()))

    for p1, p2 in pairs:
        # Passt das Paar noch auf die aktuelle Seite?
        if used_dxa + _PAIR_HEIGHT_DXA > usable_dxa:
            # Seitenumbruch einfügen
            pb = doc.add_paragraph()
            pb.paragraph_format.space_before = Pt(0)
            pb.paragraph_format.space_after  = Pt(0)
            run = pb.add_run()
            run.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK_TYPE'])
                          .WD_BREAK_TYPE.PAGE)
            used_dxa   = 0
            usable_dxa = _USABLE_REST

        _add_pair_table(doc, p1, p2)
        used_dxa += _PAIR_HEIGHT_DXA

    doc.save(output_path)
    print(f'✓ Gespeichert: {output_path}')


def _empty():
    return {k: '' for k in ['number','name','position','foot','age','height',
                              'nationality','games_display','goals','assists',
                              'yellow','yellow_red','red','transfer_info',
                              'gegentore','zu_null']} | {'local_photo': None}


if __name__ == '__main__':
    test = [
        {'number':'1','name':'Benjamin Uphoff','position':'Torwart','foot':'R',
         'age':'32','height':'1,92 m','nationality':'DEU',
         'games_display':'25 (25)','goals':'0','assists':'0',
         'yellow':'1','yellow_red':'0','red':'0','local_photo':None,
         'transfer_info':'SC Freiburg (01.07.2024), ablösefrei'},
        {'number':'30','name':'Max Hagemoser','position':'Torwart','foot':'L',
         'age':'23','height':'1,90 m','nationality':'DEU',
         'games_display':'0 (0)','goals':'0','assists':'0',
         'yellow':'0','yellow_red':'0','red':'0','local_photo':None,
         'transfer_info':'1.FC Köln U19 (01.07.2022), ablösefrei'},
    ]
    generate_etiketten(test, '/tmp/test_v4.docx', 'Test')
    print('OK')
